"""
renderer/docx.py
Renders the resume to a .docx file using python-docx.
Returns bytes so the Flask endpoint can serve it directly.

Install: pip install python-docx
"""

import io


def render_docx(data: dict) -> bytes:
    """
    Returns DOCX as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page margins ──
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── Helper: set paragraph font ──
        def para(text, bold=False, italic=False, size=12, color=None, align=None, space_before=0, space_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            if align:
                p.alignment = align
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p

        def add_hr(doc):
            """Add a thin horizontal rule."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1a1a1a")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def section_heading(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            # underline rule via paragraph border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "dddddd")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def ari_entry(e):
            # Entry title
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(e["title"])
            run.bold = True
            run.font.size = Pt(12)

            def ari_line(label, text, label_color):
                if not text:
                    return
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent  = Inches(0.2)
                p.paragraph_format.space_after  = Pt(1)
                lbl = p.add_run(f"{label}  ")
                lbl.bold = True
                lbl.font.size = Pt(10)
                lbl.font.color.rgb = RGBColor(*label_color)
                body = p.add_run(text)
                body.font.size = Pt(10)

            ari_line("Action", e["action"], (0x1e, 0x3a, 0x8a))
            ari_line("Result", e["result"], (0x14, 0x53, 0x2d))
            ari_line("Impact", e["impact"], (0x3b, 0x07, 0x64))

        def ari_section(entries, heading):
            if not entries:
                return
            section_heading(heading)
            for e in entries:
                ari_entry(e)

        # ── Name ──
        para(data["name"] or "Resume", bold=True, size=22, space_after=2)

        if data["title"]:
            para(data["title"], italic=True, size=13, color=(0x55, 0x55, 0x55), space_after=4)

        if data["purpose"]:
            para(data["purpose"], size=11, space_after=8)

        # ── Contact ──
        contact_parts = []
        if data["email"]:    contact_parts.append(data["email"])
        if data["phone"]:    contact_parts.append(data["phone"])
        if data["location"]: contact_parts.append(data["location"])
        for lnk in data["links"]:
            contact_parts.append(f"{lnk['title']}: {lnk['url']}")
        if contact_parts:
            para("  ·  ".join(contact_parts), size=10, color=(0x55, 0x55, 0x55), space_after=6)

        add_hr(doc)

        ari_section(data["experience"],   "Experience")
        ari_section(data["volunteer"],    "Volunteer Work")
        ari_section(data["certificates"], "Certificates")
        ari_section(data["awards"],       "Awards")

        if data["hobbies"]:
            section_heading("Hobbies & Interests")
            para(" · ".join(data["hobbies"]), size=11, color=(0x55, 0x55, 0x55), space_before=4)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except ImportError:
        return b""  # Caller should handle missing dependency gracefully
